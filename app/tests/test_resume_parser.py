"""Tests for resume parsing subsystem (PDF, DOCX, TXT)."""
import os
import pytest
from app.services.parser import (
    parse_resume,
    parse_resume_bytes,
    ParsedResume,
    ResumeParseError,
    UnsupportedFileTypeError,
    FileTooLargeError,
)


def test_parse_txt_bytes():
    raw_content = b"John Doe\nSoftware Engineer\nPython, FastAPI, Postgres\njohn.doe@example.com"
    result = parse_resume_bytes(raw_content, file_name="resume.txt")
    assert isinstance(result, ParsedResume)
    assert "John Doe" in result.raw_text
    assert result.file_type == "txt"
    assert result.email == "john.doe@example.com"


def test_parse_docx_bytes():
    try:
        import docx
    except ImportError:
        pytest.skip("python-docx not installed")

    import io
    doc = docx.Document()
    doc.add_heading("Jane Smith", level=1)
    doc.add_paragraph("Senior Data Scientist - PyTorch, ML, Python")
    doc.add_paragraph("Contact: jane.smith@example.org")

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    result = parse_resume_bytes(docx_bytes, file_name="jane_resume.docx")
    assert isinstance(result, ParsedResume)
    assert "Jane Smith" in result.raw_text
    assert result.file_type == "docx"
    assert result.email == "jane.smith@example.org"


def test_parse_pdf_bytes():
    pdf_text_content = b"Jane Doe Resume\nSoftware Architect\njane@example.com"
    result = parse_resume_bytes(pdf_text_content, file_name="test.pdf")
    assert isinstance(result, ParsedResume)
    assert result.file_type == "pdf"
    assert "Jane Doe" in result.raw_text
    assert result.email == "jane@example.com"


def test_parse_blank_pdf_bytes_raises_error(tmp_path):
    try:
        import pypdf
    except ImportError:
        pytest.skip("pypdf not installed")

    from pypdf import PdfWriter
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    pdf_path = str(tmp_path / "blank.pdf")
    with open(pdf_path, "wb") as f:
        writer.write(f)

    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()

    with pytest.raises(ResumeParseError, match="Could not extract text from PDF"):
        parse_resume_bytes(pdf_bytes, file_name="blank.pdf")


def test_unsupported_file_extension():
    with pytest.raises(UnsupportedFileTypeError):
        parse_resume_bytes(b"some binary content", file_name="resume.exe")


def test_file_size_exceeded():
    large_bytes = b"x" * (11 * 1024 * 1024)  # 11MB
    with pytest.raises(FileTooLargeError):
        parse_resume_bytes(large_bytes, file_name="huge.txt", max_size_bytes=10 * 1024 * 1024)


def test_corrupt_pdf_bytes_raises_parse_error():
    corrupt_pdf = b"%PDF-1.4\ncorrupted header and body data that fails parsing"
    with pytest.raises(ResumeParseError):
        parse_resume_bytes(corrupt_pdf, file_name="bad.pdf")


from app.services.parser import extract_candidate_metadata, clean_candidate_name


def test_extract_candidate_name_and_metadata():
    resume_text = (
        "SURYA PRAKASH - Senior Software Engineer\n"
        "suryaprakash@example.com | +1 555-0199\n"
        "Experienced Backend Engineer with 6+ years specializing in Python, FastAPI, and PostgreSQL."
    )
    meta = extract_candidate_metadata(resume_text, file_name="c8955ffd4d214e2eb7bdaffcf61dca52_AI-RESUME SURYA - SingleP.pdf")
    assert meta["full_name"] == "Surya Prakash"
    assert meta["email"] == "suryaprakash@example.com"
    assert meta["full_name"] != "c8955ffd4d214e2eb7bdaffcf61dca52_AI-RESUME SURYA - SingleP"


def test_clean_candidate_name_filename_cleaning():
    raw_file = "c8955ffd4d214e2eb7bdaffcf61dca52_AI-RESUME SURYA - SingleP.pdf"
    cleaned = clean_candidate_name(raw_file)
    assert cleaned == "Surya"
    assert "c8955ffd" not in cleaned
    assert "AI-RESUME" not in cleaned
    assert "SingleP" not in cleaned

    uuid_file = "550e8400-e29b-41d4-a716-446655440000_Alex_Chen_CV.pdf"
    cleaned_uuid = clean_candidate_name(uuid_file)
    assert cleaned_uuid == "Alex Chen"


def test_parsed_resume_contains_candidate_name():
    raw_content = b"SURYA PRAKASH - Senior Software Engineer\nsurya@example.com\nPython dev"
    res = parse_resume_bytes(raw_content, file_name="c8955ffd4d214e2eb7bdaffcf61dca52_AI-RESUME SURYA - SingleP.pdf")
    assert res.candidate_name == "Surya Prakash"
    assert res.email == "surya@example.com"
