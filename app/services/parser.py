"""Unified Resume Parsing Service (PDF, DOCX, TXT/MD).

Supports parsing raw bytes or file paths into a structured ParsedResume object.
Validates file extensions, size limits, and sanitizes input data.
"""
from __future__ import annotations

import io
import logging
import re
import os
from typing import Any
from pydantic import BaseModel, Field

logger = logging.getLogger("talentops.parser")

_EMAIL_REGEX = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
_ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
_DEFAULT_MAX_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB limit


class ResumeParseError(Exception):
    """Raised when parsing resume file content fails."""
    pass


class UnsupportedFileTypeError(ResumeParseError):
    """Raised when the uploaded file type is not supported."""
    pass


class FileTooLargeError(ResumeParseError):
    """Raised when file size exceeds maximum permitted limit."""
    pass


class ParsedResume(BaseModel):
    """Structured result of resume parsing."""
    raw_text: str
    file_name: str
    file_type: str
    email: str = ""
    candidate_name: str = ""
    skills: list[str] = Field(default_factory=list)
    metadata: dict[str, Any] = Field(default_factory=dict)


def clean_candidate_name(raw_name: str) -> str:
    """Clean up a raw filename or text line to extract a candidate full name."""
    if not raw_name:
        return "Candidate"

    # 1. Remove file extensions
    cleaned = re.sub(r"\.(pdf|docx|doc|txt|md)$", "", str(raw_name).strip(), flags=re.IGNORECASE)

    # 2. Remove UUID prefixes (hex 32 or standard 36-char uuid followed by _ or -)
    cleaned = re.sub(r"^[a-fA-F0-9]{32}[_-]?", "", cleaned)
    cleaned = re.sub(r"^[a-fA-F0-9]{8}-[a-fA-F0-9]{4}-[a-fA-F0-9]{4}-[a-fA-F0-9]{4}-[a-fA-F0-9]{12}[_-]?", "", cleaned)

    # 3. Replace underscores/dashes with spaces first
    cleaned = cleaned.replace("_", " ").replace("-", " ")

    # 4. Remove common junk keywords
    junk_patterns = [
        r"\bAI[ -]?RESUME\b", r"\bRESUME\b", r"\bCV\b", r"\bCURRICULUM\b", r"\bVITAE\b",
        r"\bSingle[ -]?P(age)?\b", r"\bDraft\b", r"\bFinal\b", r"\bCopy\b", r"\bUpload\b",
        r"\bDocument\b", r"\bProfile\b"
    ]
    for pat in junk_patterns:
        cleaned = re.sub(pat, "", cleaned, flags=re.IGNORECASE)

    # 5. Strip non-alpha characters and collapse whitespace
    cleaned = re.sub(r"[^a-zA-Z\s\.]", "", cleaned)
    cleaned = " ".join(cleaned.split())

    if cleaned and len(cleaned) >= 2:
        return cleaned.title()
    return "Candidate"


def extract_candidate_metadata(resume_text: str, file_name: str | None = None) -> dict[str, str]:
    """Extract candidate full name and email from top 10-15 lines of resume text."""
    email = extract_email_from_text(resume_text or "")

    extracted_name = ""
    lines = [line.strip() for line in (resume_text or "").splitlines() if line.strip()]
    top_lines = lines[:15]

    ignore_words = {
        "curriculum", "vitae", "resume", "cv", "summary", "experience", "education",
        "profile", "contact", "page", "phone", "email", "skills", "projects", "senior",
        "junior", "lead", "staff", "principal", "engineer", "developer", "architect",
        "manager", "data", "software", "fullstack", "backend", "frontend"
    }

    for line in top_lines:
        if "@" in line or "http" in line or "www." in line or "linkedin" in line or "github" in line:
            continue
        if re.search(r"\+?\d[\d\s-]{7,}", line):
            continue

        first_segment = re.split(r"\s*[-|\:\,]\s*", line)[0].strip()
        words = first_segment.split()
        if 1 <= len(words) <= 4:
            clean_words = [re.sub(r"[^a-zA-Z]", "", w) for w in words]
            clean_words = [w for w in clean_words if w]
            if clean_words and not any(w.lower() in ignore_words for w in clean_words):
                candidate_cand = " ".join(clean_words)
                if len(candidate_cand) >= 2:
                    extracted_name = candidate_cand.title()
                    break

    if not extracted_name and file_name:
        extracted_name = clean_candidate_name(file_name)
    elif not extracted_name:
        extracted_name = "Candidate"
    else:
        extracted_name = clean_candidate_name(extracted_name)

    if not email or "@" not in email:
        safe_email_name = extracted_name.lower().replace(" ", ".")
        email = f"{safe_email_name}@example.com"

    return {
        "full_name": extracted_name,
        "email": email,
    }


def extract_email_from_text(text: str) -> str:
    """Extract candidate email address from resume text."""
    matches = _EMAIL_REGEX.findall(text or "")
    if matches:
        for m in matches:
            if not any(ignore in m.lower() for ignore in ["example.com", "domain.com", "github.com", "w3.org"]):
                return m
        return matches[0]
    return ""


def parse_pdf_bytes(pdf_bytes: bytes) -> str:
    """Extract text from raw PDF bytes using pypdf with fallback to pdfplumber."""
    for pypdf_log_name in ["pypdf", "pypdf._reader", "pypdf.filters", "pypdf.generic._data_structures"]:
        logging.getLogger(pypdf_log_name).setLevel(logging.ERROR)

    is_pdf = pdf_bytes.startswith(b"%PDF") or b"%PDF-" in pdf_bytes[:1024]

    if is_pdf:
        extracted = ""
        # 1. Primary parser: pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(pdf_bytes))
            pages = [(page.extract_text() or "") for page in reader.pages]
            extracted = "\n".join(t for t in pages if t.strip()).strip()
        except Exception as e:
            logger.warning("pypdf failed to extract text from PDF: %s", e)

        # 2. Secondary fallback: pdfplumber
        if not extracted or len(extracted) < 10:
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                    pages = [(page.extract_text() or "") for page in pdf.pages]
                    extracted = "\n".join(t for t in pages if t.strip()).strip()
            except Exception as e:
                logger.warning("pdfplumber failed to extract text from PDF: %s", e)

        if not extracted or not extracted.strip():
            raise ResumeParseError("Could not extract text from PDF. Ensure PDF is not scanned/image-only.")

        return extracted

    # Fallback ONLY if file is a non-PDF plain text file
    try:
        text = pdf_bytes.decode("utf-8", errors="strict")
        if text.strip():
            logger.info("Parsed plain text file without PDF header")
            return text
    except Exception:
        pass

    raise ResumeParseError("Failed to parse PDF content: Invalid or corrupt PDF binary structure")


def parse_docx_bytes(docx_bytes: bytes) -> str:
    """Extract text from raw DOCX bytes using python-docx or zipfile XML fallback."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(docx_bytes))
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += "\n" + row_text
        return text
    except ImportError:
        logger.warning("python-docx not installed; attempting XML extraction fallback")
        try:
            import zipfile
            from xml.etree import ElementTree as ET
            with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
                xml_content = zf.read("word/document.xml")
                tree = ET.fromstring(xml_content)
                paragraphs = []
                for p in tree.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
                    texts = [t.text for t in p.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t") if t.text]
                    if texts:
                        paragraphs.append("".join(texts))
                return "\n".join(paragraphs)
        except Exception as ex:
            raise ResumeParseError(f"Failed to parse DOCX bytes via XML fallback: {ex}") from ex
    except Exception as e:
        logger.error("Failed to parse DOCX bytes: %s", e)
        raise ResumeParseError(f"Failed to parse DOCX content: {e}") from e


def parse_resume_bytes(
    content: bytes,
    file_name: str = "resume.pdf",
    max_size_bytes: int = _DEFAULT_MAX_SIZE_BYTES,
) -> ParsedResume:
    """Parse resume raw bytes into a ParsedResume object with strict validation."""
    if len(content) > max_size_bytes:
        raise FileTooLargeError(
            f"File size ({len(content)} bytes) exceeds maximum limit of {max_size_bytes} bytes"
        )

    ext = os.path.splitext(file_name)[1].lower() or ".pdf"
    if ext not in _ALLOWED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"File extension '{ext}' is not supported. Supported extensions: {', '.join(sorted(_ALLOWED_EXTENSIONS))}"
        )

    file_type = ext.lstrip(".")

    if file_type == "pdf":
        raw_text = parse_pdf_bytes(content)
    elif file_type == "docx":
        raw_text = parse_docx_bytes(content)
    else:  # txt or md
        try:
            raw_text = content.decode("utf-8", errors="replace")
        except Exception as e:
            raise ResumeParseError(f"Failed to decode text file: {e}") from e

    if not raw_text or not raw_text.strip():
        raise ResumeParseError("Extracted text from resume is empty or missing")

    meta = extract_candidate_metadata(raw_text, file_name=file_name)
    email = meta.get("email") or extract_email_from_text(raw_text)
    candidate_name = meta.get("full_name") or clean_candidate_name(file_name)

    return ParsedResume(
        raw_text=raw_text,
        file_name=file_name,
        file_type=file_type,
        email=email,
        candidate_name=candidate_name,
        metadata={"content_length": len(content), "char_count": len(raw_text)}
    )


def parse_resume(path: str) -> ParsedResume:
    """Parse a resume file path."""
    if not os.path.exists(path):
        raise ResumeParseError(f"Resume file path does not exist: {path}")

    with open(path, "rb") as f:
        content = f.read()

    filename = os.path.basename(path)
    return parse_resume_bytes(content, file_name=filename)
